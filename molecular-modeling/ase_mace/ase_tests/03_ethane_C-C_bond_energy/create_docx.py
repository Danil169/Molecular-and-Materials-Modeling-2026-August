from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Отчет по расчету энергии связи C-H в метане (Challenge I.4)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Theory
doc.add_heading('1. Теоретическая база и методология', level=1)
doc.add_paragraph(
    "Для вычисления энергии разрыва связи (BDE) C-H в метане моделируется реакция гомолитического разрыва:\n"
    "CH4 → CH3• + H•"
)
doc.add_paragraph(
    "Энергия связи вычисляется как разность полных энергий изолированных продуктов и исходной молекулы:\n"
    "BDE(C-H) = E(CH3•) + E(H•) - E(CH4)"
)
doc.add_paragraph(
    "С помощью библиотеки ASE (Atomic Simulation Environment) была построена геометрия молекулы метана и метильного радикала. "
    "Структуры были оптимизированы (релаксированы) алгоритмом BFGS. "
    "Расчеты проводились с использованием двух методов:"
)
doc.add_paragraph("1. EMT (Effective Medium Theory) – классический эмпирический потенциал.", style='List Bullet')
doc.add_paragraph("2. MACE – современный нейросетевой потенциал машинного обучения.", style='List Bullet')

# 2. Results Table
doc.add_heading('2. Результаты расчетов', level=1)
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Параметр'
hdr_cells[1].text = 'EMT'
hdr_cells[2].text = 'MACE'
hdr_cells[3].text = 'Эксперимент'
hdr_cells[4].text = 'Ошибка EMT'
hdr_cells[5].text = 'Ошибка MACE'

# Data rows
records = (
    ('Длина связи C-H', '1.1515 Å', '1.0931 Å', '1.0870 Å', '5.93%', '0.56%'),
    ('Энергия разрыва (эВ)', '3.2539 эВ', '4.3158 эВ', '4.55 эВ', '-1.29 эВ', '-0.23 эВ'),
    ('Энергия разрыва (кДж/моль)', '313.95', '416.41', '439.0', '28.49%', '5.15%')
)
for item in records:
    row_cells = table.add_row().cells
    for i in range(6):
        row_cells[i].text = item[i]

# 3. Conclusions
doc.add_heading('3. Выводы', level=1)
p1 = doc.add_paragraph()
p1.add_run('Геометрия (Длина связи): ').bold = True
p1.add_run(
    "Нейросеть MACE практически идеально предсказала геометрию молекулы метана. "
    "Длина связи отличается от экспериментальной всего на 0.006 Å (ошибка 0.5%). "
    "Классический потенциал EMT сильно завысил длину связи, растянув её до 1.15 Å."
)

p2 = doc.add_paragraph()
p2.add_run('Энергия связи (BDE): ').bold = True
p2.add_run(
    "Связь C-H в метане очень прочная (439 кДж/моль). "
    "Потенциал EMT сильно занизил эту энергию (ошибка 28.5%), так как он не умеет правильно описывать "
    "электронную структуру sp3-гибридизованного углерода и ковалентные органические связи. "
    "Потенциал MACE дал значение 416.4 кДж/моль, показав высокую точность (ошибка около 5%), "
    "сравнимую с ресурсоемкими квантово-химическими методами (DFT/CCSD), выполнив расчет за доли секунды."
)

# Save
file_path = 'Methane_CH_Bond_Report.docx'
doc.save(file_path)
print(f"File saved to {file_path}")
